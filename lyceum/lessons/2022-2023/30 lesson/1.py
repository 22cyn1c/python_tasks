import docx
import sys
document = docx.Document()

place = input()
time = input()
a = []
for i in sys.stdin:
    a.append(i)
for i in a:
    document.add_heading(f'Дорогая {i}!')

    document.add_paragraph(f'Поздравляю вас с 8 марта и приглашаю тебя на мероприятие {place}.')

    p = document.add_paragraph('Приходите в ')
    p.add_run(f'{time}.').bold = True

    document.add_page_break()


document.save('test.docx')